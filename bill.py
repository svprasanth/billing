import tkinter as tk
from tkinter import filedialog, messagebox
import pyodbc
import gtts
import os
import pyttsx3
from datetime import date
import threading
import docx
from docx import Document
from num2words import num2words
import word2number as w2n
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import locale
import inflect
import sqlite3
import tkinter as ttk
import pygame
p= inflect.engine()
root = tk.Tk()
def play_mp3(file_path):
    pygame.mixer.init()
    pygame.mixer.music.load(file_path)
    pygame.mixer.music.play()

def speak(text):
    engine = pyttsx3.init()
    rate = engine.getProperty('rate')  # Get the current speaking rate (words per minute)
    engine.setProperty('rate', rate - 50)  # Decrease the speaking rate by 50 words per minute

    voices = engine.getProperty('voices')  # Get available voices
    female_voice = None
    for voice in voices:
        if "female" in voice.name.lower():
            female_voice = voice
            break

    if female_voice:
        engine.setProperty('voice', female_voice.id)  # Set the voice to a female voice if available

    engine.say(text)
    engine.runAndWait()


total_values=[]
cgst_values=[]
sgst_values=[]
def create_widgets():
    speak(text="welcome to universal  clean care product. Kindly  place your orders here")
    play_mp3("car.mp3")


    global cart_listbox
    global Date_entry
    global totaL_entry
    customer_name = tk.Label(root, text='CUSTOMER NAME')
    customer_name.grid(row=0, column=0)
    sub_total_label = tk.Label(root, text='TOTAL AFTER TAX')
    sub_total_label.grid(row=0, column=2)
    total_entry = tk.Entry(root, textvariable=TOTAL)
    total_entry.grid(row=0, column=3)
    customer_entry = tk.Entry(root, textvariable=customer)
    customer_entry.grid(row=0, column=1)
    customer_address_label = tk.Label(root, text='CUSTOMER ADDRESS')
    customer_address_label.grid(row=1, column=0)
    customer_address_entry = tk.Entry(root, textvariable=customer_address, width=60)
    customer_address_entry.grid(row=1, column=1)
    Date=tk.Label(root,text='Invoice Date')
    Date.grid(row=3,column=2)
    invoicenumber = tk.Label(root, text='Invoice Number')
    invoicenumber.grid(row=2, column=2)
    invoicenumber_entry = tk.Entry(root, textvariable=INVOICE)
    invoicenumber_entry.grid(row=2, column=3)
    Date_entry = tk.Entry(root)
    Date_entry.grid(row=3, column=3)
    Date_entry.insert(0,date.today().strftime('%d-%m-%y'))
    customer_GST_label = tk.Label(root, text='CUSTOMER GST NUMBER')
    customer_GST_label.grid(row=2, column=0)
    customer_GST_entry = tk.Entry(root, textvariable=customer_gstin_number, width=60)
    customer_GST_entry.grid(row=2, column=1)
    PRODUCT_NAME = tk.Label(root, text='PRODUCT_NAME')
    PRODUCT_NAME.grid(row=3, column=0)
    PRICE_label = tk.Label(root, text='QUANTITY')
    PRICE_label.grid(row=4, column=0)
    customer_GST_label = tk.Label(root, text='PRICE')
    customer_GST_label.grid(row=5, column=0)


    product_price = tk.Entry(root, textvariable=price_of_product)
    product_price.grid(row=5, column=1)
    product_quantity = tk.Entry(root, textvariable=quantity_of_product)
    product_quantity.grid(row=4, column=1)
    product_name = tk.Entry(root, textvariable=name_of_product)
    product_name.grid(row=3, column=1)
    customer_GST_label = tk.Label(root, text='your cart', font='arialblack', fg='#F753A7')
    customer_GST_label.grid(row=9, column=1)

    cart_listbox = tk.Listbox(root, selectmode=tk.MULTIPLE, width=100, height=15)
    cart_listbox.grid(row=10, column=2)
    total_button = tk.Button(root, text='total', font='GThaptik', background='grey', command=update_sub_total)
    total_button.grid(row=9, column=2)
    total_button = tk.Button(root, text='reset', font='GThaptik', background='grey', command=reset_values)
    total_button.grid(row=10, column=3)
    total_button = tk.Button(root, text='invoice', font='GThaptik', background='grey', command=generate_invoice)
    total_button.grid(row=10, column=4)
    customer_GST_label = tk.Label(root, text='Rupees in words', font='arialblack', fg='#F753A7')
    customer_GST_label.grid(row=12, column=0)

    cgst_sgst_label = tk.Label(root, text='Tax Rate (in %)')
    cgst_sgst_label.grid(row=1, column=2)
    tax_rate_entry = tk.Entry(root, textvariable=tax_rate_input)
    tax_rate_entry.grid(row=1, column=3)
    view_button = tk.Button(root, text="View Database", font="GThaptik", background="grey", command=view_database)
    view_button.grid(row=13, column=5)

    def add_to_cart():
        product_name = name_of_product.get()
        price = float(price_of_product.get())
        quantity = int(quantity_of_product.get())
        total = price * quantity

        if product_name and price > 0 and quantity > 0:
            tax_rate_value = float(tax_rate_input.get())
            cgst, sgst, total_with_tax = calculate_cgst_sgst_total(total, tax_rate_value)
            cart_item = f"{product_name} |  ₹{price:.2f} |  {quantity} |  ₹{total:.2f} | CGST: ₹{cgst:.2f} | SGST: ₹{sgst:.2f} | Total After Tax: ₹{total_with_tax:.2f}"
            cart_listbox.insert(tk.END, cart_item)
            total_values.append(total)
            cgst_values.append(cgst)
            sgst_values.append(sgst)

            # Clear input fields
            name_of_product.set("")
            price_of_product.set("")
            quantity_of_product.set("")

    cart_button = tk.Button(root, text='Add to cart', font='GThaptik', background='grey', command=add_to_cart)
    cart_button.grid(row=9, column=0)


def calculate_total_sum():
    total_sum = sum(total_values)
    return total_sum
def cgst():
    total_sum = sum(total_values)
    cgst_tax= sum(cgst_values)
    return cgst_tax

def sgst():
    total_sum = sum(total_values)
    sgst_tax= sum(sgst_values)
    return sgst_tax
def create_summary_table(document):
    summary_table = document.add_table(rows=5, cols=4)
    summary_table.style = 'Table Grid'
    summary_table.cell(0, 0).text = 'Subtotal:'
    summary_table.cell(0, 3).text = f"₹{calculate_total_sum():.2f}"
    summary_table.cell(1, 0).text = 'CGST:'
    summary_table.cell(1, 3).text =f"₹{cgst():.0f}"
    summary_table.cell(2, 0).text = 'SGST:'
    summary_table.cell(2, 3).text = f"₹{sgst():.0f}"
    total_after_tax_value=round(update_sub_total(),0)
    locale.setlocale(locale.LC_ALL, 'en_IN')
    total_after_tax_words = p.number_to_words(int(total_after_tax_value)).replace(',', ' ')
    summary_table.cell(3, 0).text = 'Total After Tax:'
    summary_table.cell(3, 3).text = f"₹{round( update_sub_total(),0):.0f}"
    summary_table.cell(4,0).text='Total In Words : '
    summary_table.cell(4,3).text=f"{total_after_tax_words.capitalize()} Rupees Only"
    for row in summary_table.rows:
        row.cells[3].merge(row.cells[2])





def generate_document(customer_name, customer_address, invoice_number, invoice_date, cart_items, gst_no,total_before_tax,inclusive_of_tax,cgst_total,sgst_total):
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
    heading = "Tax Invoice"
    document = Document()
    PARAGRAPH=document.add_heading(heading,level=1)
    PARAGRAPH.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER

    # Create and format the customer details table
    document.add_paragraph("Universal Clean Care Products,\n15/16 vivekanandhapuram(2nd Cross Street),\nBeemanagar,\n Trichy - 620001.\n GSTIN: 33BPXPV8558P1ZC\n MOBILE NO: 7667367316, 9786794916")

    customer_table = document.add_table(rows=5, cols=2)  # Increase the number of rows
    customer_table.style = 'Table Grid'
    customer_table.cell(2, 0).text = 'Customer Name:'
    customer_table.cell(2, 1).text = customer_name
    customer_table.cell(1, 0).text = 'Invoice Date'
    customer_table.cell(1,1).text=invoice_date
    customer_table.cell(3,0).text= 'Customer Address'
    customer_table.cell(3, 1).text = customer_address
    customer_table.cell(0, 0).text = 'Invoice Number:'
    customer_table.cell(0, 1).text = invoice_number
    customer_table.cell(4, 0).text = 'Customer GST NUMBER:'
    customer_table.cell(4, 1).text = gst_no

    # Add Total Before Tax, CGST, SGST, Total After Tax



    for row in customer_table.rows:
        for cell in row.cells:
            cell.width = docx.shared.Inches(5)  # Adjust the width as needed
            cell.height = docx.shared.Inches(0.5)  # Adjust the height as needed

    # Add a newline for spacing
    document.add_paragraph()

    # Create and format the invoice details table
    invoice_table = document.add_table(rows=len(cart_items) + 1, cols=4)
    invoice_table.style = 'Table Grid'
    invoice_table.cell(0, 0).text = 'Product Name'
    invoice_table.cell(0, 1).text = 'Price'
    invoice_table.cell(0, 2).text = 'Quantity'
    invoice_table.cell(0, 3).text = 'Total'

    # Populate the invoice details table
    for idx, cart_item in enumerate(cart_items, start=1):
        product_name, price, quantity, total = cart_item.split(' | ')[0:4]
        invoice_table.cell(idx, 0).text = product_name
        invoice_table.cell(idx, 1).text = price
        invoice_table.cell(idx, 2).text = quantity
        invoice_table.cell(idx, 3).text = total

    create_summary_table(document)
    document.add_paragraph("\n\nfor UNIVERSAL CLEAN CARE PRODUCTS\n\n\n\n Properitor").alignment= WD_PARAGRAPH_ALIGNMENT.RIGHT
    # Save the document
    document.save(file_path)

def create_invoice_database_table():
    conn = sqlite3.connect('invoice_database.db')
    cursor = conn.cursor()

    create_table = '''
    CREATE TABLE IF NOT EXISTS invoices (
        id INTEGER PRIMARY KEY,
        customer_name TEXT,
        customer_address TEXT,
        invoice_number TEXT,
        invoice_date TEXT,
        products TEXT,
        gst_no TEXT,
        total_before_tax REAL,
        total_after_tax REAL,
        tax_rate REAL
    )
    '''
    cursor.execute(create_table)

    conn.commit()
    conn.close()

def generate_invoice():
    create_invoice_database_table()
    customer_name_value = customer.get()
    customer_address_value = customer_address.get()
    invoice_number_value = INVOICE.get()
    invoice_date_value = Date_entry.get()
    cart_items = cart_listbox.get(0, tk.END)
    gst_no_value=customer_gstin_number.get()
    cgst_rate_tax_value= int(tax_rate_input.get())*0.5
    sgst_rate_tax_value= int(tax_rate_input.get())*0.5
    total_before_tax= calculate_total_sum()
    inclusive_of_tax=update_sub_total()
    rupees_in_words=str(rupees_in_wordss.get())
    invoice_date_value = Date_entry.get()
    generate_document(customer_name_value, customer_address_value, invoice_number_value, invoice_date_value, cart_items,gst_no_value,cgst_rate_tax_value,sgst_rate_tax_value,calculate_total_sum,update_sub_total)
    conn = sqlite3.connect('invoice_database.db')
    cursor = conn.cursor()
    insert_query = '''
            INSERT INTO invoices (
                customer_name, customer_address, invoice_number, invoice_date,
                products, gst_no, total_before_tax, total_after_tax, tax_rate
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        '''

    products_str = '\n'.join(cart_items)  # Combine the cart items into a single string

    values = (
        customer_name_value, customer_address_value, invoice_number_value,
        invoice_date_value, products_str, gst_no_value, total_before_tax,
        inclusive_of_tax, float(tax_rate_input.get())
    )
    cursor.execute(insert_query, values)
    conn.commit()
    conn.close()
    speak("Invoice generated successfully.")

def view_database():
    # Create a new window
    view_window = tk.Toplevel(root)
    view_window.title("View Database")

    # Create a Treeview widget to display the data
    listbox = tk.Listbox(view_window, width=800, height=20)
    listbox.pack()

    # Retrieve data from the database and populate the Listbox
    conn = sqlite3.connect('invoice_database.db')
    cursor = conn.cursor()
    query ="SELECT customer_name, invoice_number, customer_address, gst_no, products, total_before_tax, total_after_tax, tax_rate,invoice_date FROM invoices"
    cursor.execute(query)
    rows = cursor.fetchall()

    for row in rows:
        listbox.insert(tk.END,f"Customer: {row[0]}\n    Invoice: {row[1]}\n Address: {row[2]}\n GSTIN: {row[3]}\n   Products:\n{row[4]}\n   Total Before Tax: ₹{row[5]:.2f} \nTotal After Tax: ₹{row[6]:.2f}    \nTax Rate: {row[7]} \nInvoice_date: ₹{row[8]}")

    conn.close()

    # Retrieve data from the database and populate the Treeview
    conn = sqlite3.connect('invoice_database.db')
    cursor = conn.cursor()

    query = "SELECT * FROM invoices"
    cursor.execute(query)
    rows = cursor.fetchall()

    for row in rows:
        listbox.insert(tk.END, f"Customer: {row[0]}, Invoice: {row[1]}, Total After Tax: ₹{row[2]:.2f}")

    conn.close()



def update_sub_total():
    total = sum(float(item.split('| Total After Tax: ₹')[1]) for item in cart_listbox.get(0, tk.END))
    TOTAL.set(f"₹{total:.2f}")
    return total

def reset_values():
    speak("have a nice day. Thank You for shopping with us!. kindly visit again")
    name_of_product.set("")
    price_of_product.set("")
    quantity_of_product.set("")
    customer.set("")
    customer_address.set("")
    customer_gstin_number.set("")
    TOTAL.set("0.00")
    cart_listbox.delete(0, tk.END)

def calculate_cgst_sgst_total(total, tax_rate):
    cgst = (total * tax_rate) / 200
    sgst = (total * tax_rate) / 200
    total_with_tax = total + cgst + sgst
    return cgst, sgst, total_with_tax





root.geometry("1900x1900")
customer = tk.StringVar()
customer_address = tk.StringVar()
customer_gstin_number = tk.StringVar()
price_of_product = tk.StringVar()
quantity_of_product = tk.StringVar()
tax_rate = tk.StringVar()
TOTAL = tk.StringVar()
converted_num_to_word=tk.StringVar()
INVOICE= tk.StringVar()
rupees_in_wordss= tk.StringVar()
total_after_tax = tk.StringVar()
name_of_product = tk.StringVar()
tax_rate_input = tk.StringVar()
Date1=date.today()
create_widgets()
root.title("Billing Software")
root.mainloop()
